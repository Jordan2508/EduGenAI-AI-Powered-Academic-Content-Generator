from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import io


def create_docx(text):
    doc = Document()
    doc.add_heading("EduGenAI Generated Content", level=1)
    doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def create_pdf(text):
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph("<b>EduGenAI Generated Content</b>", styles["Heading1"]))
    story.append(Paragraph(text.replace("\n", "<br/>"), styles["BodyText"]))

    doc.build(story)

    buffer.seek(0)

    return buffer